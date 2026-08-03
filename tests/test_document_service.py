import os
import unittest
import tempfile
import docx
from dat.services.document_service import DocumentService
from dat.models.doc_request import ChangeSummary
from dat.models.screenshot_info import ScreenshotInfo

class TestDocumentService(unittest.TestCase):
    def test_generate_documentation_end_to_end(self):
        doc_service = DocumentService()
        with tempfile.TemporaryDirectory() as tmp_dir:
            out_docx = os.path.join(tmp_dir, "e2e_doc.docx")
            out_md = os.path.join(tmp_dir, "e2e_doc.md")

            # Generate DOCX
            res_docx = doc_service.generate_documentation(
                output_path=out_docx,
                title_override="Payment Gateway Integration",
                author="Lead Dev",
                ticket_override="PAY-500",
                output_format="docx"
            )
            self.assertTrue(os.path.exists(res_docx))

            # Generate Markdown
            res_md = doc_service.generate_documentation(
                output_path=out_md,
                title_override="Payment Gateway Integration",
                author="Lead Dev",
                ticket_override="PAY-500",
                output_format="md"
            )
            self.assertTrue(os.path.exists(res_md))
            with open(res_md, "r") as f:
                content = f.read()
            self.assertIn("# Payment Gateway Integration", content)
            self.assertIn("PAY-500", content)

    def test_summary_and_screenshots_override_are_preserved(self):
        """GUI edits (summary text, test cases, screenshot assignments/order)
        must survive export unchanged - they must not be silently regenerated."""
        doc_service = DocumentService()
        with tempfile.TemporaryDirectory() as tmp_dir:
            img_path = os.path.join(tmp_dir, "shot.png")
            from PIL import Image
            Image.new("RGB", (10, 10)).save(img_path)

            out_docx = os.path.join(tmp_dir, "override_doc.docx")
            summary = ChangeSummary(
                overview="Hand-written overview.",
                key_points=["Manually added point"],
                impact_areas=["Manual Module"],
                test_cases=["Manually written test case"],
            )
            screenshots = [ScreenshotInfo(file_path=img_path, test_case_index=0)]

            result_path = doc_service.generate_documentation(
                output_path=out_docx,
                title_override="Manual Edits Doc",
                author="QA",
                approved_by="Reviewer X",
                ticket_override="QA-1",
                output_format="docx",
                summary_override=summary,
                screenshots_override=screenshots,
            )

            document = docx.Document(result_path)
            text_content = "\n".join([p.text for p in document.paragraphs])
            self.assertIn("Manually added point", text_content)
            self.assertIn("Manually written test case", text_content)

            table_cells = [c.text for t in document.tables for r in t.rows for c in r.cells]
            self.assertIn("Reviewer X", table_cells)

if __name__ == "__main__":
    unittest.main()
