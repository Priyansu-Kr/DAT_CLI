import os
import shutil
import tempfile
import unittest

import docx
from PIL import Image

from dat.models.screenshot_info import ScreenshotInfo
from dat.models.template_model import (
    BLOCK_BULLET_LIST,
    BLOCK_CODE,
    BLOCK_HEADING,
    BLOCK_IMAGE,
    BLOCK_PARAGRAPH,
    BLOCK_SCREENSHOTS,
    BLOCK_SEPARATOR,
    BLOCK_SUBHEADING,
    BLOCK_TABLE,
    BLOCK_TWO_COLUMNS,
    DocumentTemplate,
    TemplateBlock,
    TemplateContext,
    TemplateSection,
)
from dat.renderers.template_docx_renderer import TemplateDocxRenderer


class TestTemplateDocxRenderer(unittest.TestCase):
    def setUp(self):
        self.tmp_dir = tempfile.mkdtemp(prefix="dat-docx-")
        self.renderer = TemplateDocxRenderer()

    def tearDown(self):
        shutil.rmtree(self.tmp_dir, ignore_errors=True)

    def _out(self, name="out.docx") -> str:
        return os.path.join(self.tmp_dir, name)

    def _make_image(self, name: str, size=(60, 120)) -> str:
        path = os.path.join(self.tmp_dir, name)
        Image.new("RGB", size, color="blue").save(path)
        return path

    def _render(self, template, context=None, **kwargs) -> str:
        return self.renderer.render(
            template, context or TemplateContext(), self._out(), **kwargs
        )

    def test_renders_every_block_kind_without_error(self):
        section = TemplateSection(title="Everything", blocks=[
            TemplateBlock(kind=BLOCK_HEADING, text="Heading", level=1),
            TemplateBlock(kind=BLOCK_SUBHEADING, text="Sub", level=2),
            TemplateBlock(kind=BLOCK_PARAGRAPH, text="Body text"),
            TemplateBlock(kind=BLOCK_BULLET_LIST, items=["one", "two"]),
            TemplateBlock(kind=BLOCK_TABLE, table_headers=["A", "B"], table_rows=[["1", "2"]]),
            TemplateBlock(kind=BLOCK_IMAGE, image_path=self._make_image("i.png"), caption="Fig 1"),
            TemplateBlock(kind=BLOCK_CODE, text="print('hi')", language="python"),
            TemplateBlock(kind=BLOCK_TWO_COLUMNS, columns=["left", "right"]),
            TemplateBlock(kind=BLOCK_SEPARATOR),
        ])
        path = self._render(DocumentTemplate(name="T", sections=[section]))

        self.assertTrue(os.path.exists(path))
        document = docx.Document(path)
        texts = [p.text for p in document.paragraphs]
        self.assertIn("Everything", texts)
        self.assertIn("Heading", texts)
        self.assertIn("Sub", texts)
        self.assertIn("Body text", texts)
        self.assertIn("one", texts)
        self.assertIn("Fig 1", texts)
        # table + code block + two-column layout
        self.assertEqual(len(document.tables), 3)
        self.assertEqual(document.tables[0].cell(0, 0).text, "A")
        self.assertEqual(document.tables[0].cell(1, 1).text, "2")

    def test_disabled_section_is_not_exported(self):
        template = DocumentTemplate(name="T", sections=[
            TemplateSection(title="Kept", blocks=[TemplateBlock(kind=BLOCK_PARAGRAPH, text="kept")]),
            TemplateSection(
                title="Dropped", enabled=False,
                blocks=[TemplateBlock(kind=BLOCK_PARAGRAPH, text="dropped")],
            ),
        ])
        document = docx.Document(self._render(template))
        texts = " ".join(p.text for p in document.paragraphs)
        self.assertIn("kept", texts)
        self.assertNotIn("dropped", texts)

    def test_section_overrides_take_priority(self):
        section = TemplateSection(title="S", blocks=[TemplateBlock(kind=BLOCK_PARAGRAPH, text="body")])
        template = DocumentTemplate(name="T", sections=[section])
        document = docx.Document(
            self._render(template, section_overrides={section.section_id: False})
        )
        self.assertNotIn("body", " ".join(p.text for p in document.paragraphs))

    def test_tokens_are_resolved(self):
        template = DocumentTemplate(name="T", sections=[TemplateSection(
            title="{{ticket_id}}",
            blocks=[TemplateBlock(kind=BLOCK_PARAGRAPH, text="Written by {{author}}")],
        )])
        context = TemplateContext(ticket_id="ABC-7", author="Priyansu")
        document = docx.Document(self._render(template, context))

        texts = [p.text for p in document.paragraphs]
        self.assertIn("ABC-7", texts)
        self.assertIn("Written by Priyansu", texts)

    def test_headers_omitted_when_include_headers_off(self):
        block = TemplateBlock(
            kind=BLOCK_TABLE, table_headers=["A", "B"], table_rows=[["1", "2"]], include_headers=False
        )
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        table = docx.Document(self._render(template)).tables[0]

        self.assertEqual(len(table.rows), 1)
        self.assertEqual(table.cell(0, 0).text, "1")

    def test_many_added_rows_all_export(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["Case", "Status"])
        block.set_table_size(0, 2)
        for index in range(12):
            block.add_row()
            block.set_cell(index, 0, f"case {index}")
            block.set_cell(index, 1, "Pass")
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        table = docx.Document(self._render(template)).tables[0]

        self.assertEqual(len(table.rows), 13)  # header + 12 content rows
        self.assertEqual(table.cell(12, 0).text, "case 11")

    def test_table_with_headers_but_no_rows_exports_header_only(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["A", "B"], table_rows=[])
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        document = docx.Document(self._render(template))

        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.tables[0].rows), 1)
        self.assertEqual(document.tables[0].cell(0, 1).text, "B")

    def test_empty_table_without_headers_is_skipped(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["A"], table_rows=[], include_headers=False)
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        self.assertEqual(len(docx.Document(self._render(template)).tables), 0)

    def test_column_widths_follow_the_weights(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["Index", "Case", "Status"])
        block.set_table_size(1, 3)
        block.set_col_weight(0, 1)
        block.set_col_weight(1, 4)
        block.set_col_weight(2, 1)
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        table = docx.Document(self._render(template)).tables[0]

        widths = [column.width for column in table.columns]
        self.assertTrue(all(w is not None for w in widths), widths)
        # Middle column is four times either neighbour, and the row fills the page.
        self.assertAlmostEqual(widths[1] / widths[0], 4.0, delta=0.05)
        self.assertAlmostEqual(widths[0], widths[2], delta=1000)
        self.assertAlmostEqual(sum(w.inches for w in widths), 6.5, delta=0.05)
        # Every cell repeats its column's width, which is what Word reads.
        self.assertAlmostEqual(table.cell(0, 1).width / table.cell(0, 0).width, 4.0, delta=0.05)

    def test_equal_weights_split_the_page_evenly(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["A", "B"])
        block.set_table_size(1, 2)
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        widths = [c.width.inches for c in docx.Document(self._render(template)).tables[0].columns]
        self.assertAlmostEqual(widths[0], widths[1], places=3)

    def test_table_layout_is_fixed_so_word_keeps_the_widths(self):
        from docx.oxml.ns import qn

        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["A", "B"])
        block.set_table_size(1, 2)
        block.set_col_weight(1, 3)
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        table = docx.Document(self._render(template)).tables[0]

        layouts = table._tbl.tblPr.findall(qn("w:tblLayout"))
        self.assertEqual(len(layouts), 1)
        self.assertEqual(layouts[0].get(qn("w:type")), "fixed")

    def test_ragged_table_row_does_not_crash(self):
        block = TemplateBlock(kind=BLOCK_TABLE, table_headers=["A", "B", "C"], table_rows=[["1"]])
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        table = docx.Document(self._render(template)).tables[0]

        self.assertEqual(len(table.columns), 3)
        self.assertEqual(table.cell(1, 2).text, "")

    def test_missing_image_path_is_skipped(self):
        block = TemplateBlock(kind=BLOCK_IMAGE, image_path="/definitely/not/here.png", caption="ghost")
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[block])])
        document = docx.Document(self._render(template))
        self.assertNotIn("ghost", " ".join(p.text for p in document.paragraphs))

    def test_corrupt_image_does_not_abort_export(self):
        bad_path = os.path.join(self.tmp_dir, "bad.png")
        with open(bad_path, "w", encoding="utf-8") as f:
            f.write("not really a png")
        template = DocumentTemplate(name="T", sections=[TemplateSection(title="S", blocks=[
            TemplateBlock(kind=BLOCK_IMAGE, image_path=bad_path, caption="broken"),
            TemplateBlock(kind=BLOCK_PARAGRAPH, text="still here"),
        ])])
        document = docx.Document(self._render(template))
        self.assertIn("still here", " ".join(p.text for p in document.paragraphs))

    def test_screenshots_block_groups_and_pairs_portraits(self):
        shots = [
            ScreenshotInfo(file_path=self._make_image("p1.png", (60, 120))),
            ScreenshotInfo(file_path=self._make_image("p2.png", (60, 120))),
            ScreenshotInfo(file_path=self._make_image("w1.png", (200, 80))),
        ]
        template = DocumentTemplate(name="T", sections=[TemplateSection(
            title="S", show_title=False,
            blocks=[TemplateBlock(kind=BLOCK_SCREENSHOTS, text="Evidence")],
        )])
        document = docx.Document(
            self._render(template, TemplateContext(test_cases=["Case A"]), screenshots=shots)
        )

        texts = [p.text for p in document.paragraphs]
        self.assertIn("Evidence", texts)
        self.assertIn("Test Case 1 : Case A", texts)
        # The two portrait shots share one 2-cell layout table.
        self.assertEqual(len(document.tables), 1)
        self.assertEqual(len(document.tables[0].columns), 2)

    def test_screenshots_block_without_screenshots_renders_nothing(self):
        template = DocumentTemplate(name="T", sections=[TemplateSection(
            title="S", show_title=False, blocks=[TemplateBlock(kind=BLOCK_SCREENSHOTS, text="Evidence")],
        )])
        document = docx.Document(self._render(template))
        self.assertNotIn("Evidence", " ".join(p.text for p in document.paragraphs))

    def test_unknown_block_kind_is_ignored(self):
        section = TemplateSection(title="S", blocks=[
            TemplateBlock(kind="future_widget", text="ignored"),
            TemplateBlock(kind=BLOCK_PARAGRAPH, text="kept"),
        ])
        document = docx.Document(self._render(DocumentTemplate(name="T", sections=[section])))
        texts = " ".join(p.text for p in document.paragraphs)
        self.assertIn("kept", texts)
        self.assertNotIn("ignored", texts)

    def test_multiline_paragraph_keeps_one_paragraph(self):
        block = TemplateBlock(kind=BLOCK_PARAGRAPH, text="line one\nline two")
        template = DocumentTemplate(name="T", sections=[
            TemplateSection(title="S", show_title=False, blocks=[block])
        ])
        document = docx.Document(self._render(template))
        matching = [p for p in document.paragraphs if "line one" in p.text]
        self.assertEqual(len(matching), 1)
        self.assertIn("line two", matching[0].text)

    def test_creates_missing_output_directory(self):
        nested = os.path.join(self.tmp_dir, "a", "b", "doc.docx")
        path = self.renderer.render(DocumentTemplate.starter(), TemplateContext(), nested)
        self.assertTrue(os.path.exists(path))

    def test_empty_template_still_produces_a_file(self):
        path = self._render(DocumentTemplate(name="Empty", sections=[]))
        self.assertTrue(os.path.exists(path))


if __name__ == "__main__":
    unittest.main()
