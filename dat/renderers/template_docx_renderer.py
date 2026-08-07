"""Renders a user-authored :class:`DocumentTemplate` to a .docx file.

Deliberately mirrors, block for block, the mapping used by
``dat.gui.state.build_template_preview_content`` so the exported document
matches the live preview. Visual conventions (Arial, black headings, 1"
margins, Table Grid) follow DocxRenderer to keep both outputs consistent.
"""
import os
from typing import Dict, List, Optional

import docx
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image

from dat.models.screenshot_info import ScreenshotInfo
from dat.models.template_model import (
    BLOCK_BULLET_LIST,
    BLOCK_CODE,
    BLOCK_HEADING,
    BLOCK_IMAGE,
    BLOCK_PARAGRAPH,
    BLOCK_SCREENSHOTS,
    BLOCK_SEPARATOR,
    BLOCK_SUBHEADING,
    BLOCK_TABLE,
    BLOCK_TWO_COLUMNS,
    DocumentTemplate,
    TemplateBlock,
    TemplateContext,
)
from dat.renderers.screenshot_grouping import group_screenshots_by_test_case

FONT_NAME = "Arial"
CODE_FONT_NAME = "Courier New"
BODY_SIZE_PT = 11
HEADING_SIZES_PT = {1: 16, 2: 13, 3: 12}
CONTENT_WIDTH_IN = 6.5
MAX_PORTRAIT_WIDTH_IN = 3.0
MAX_IMAGE_HEIGHT_IN = 7.0


class TemplateDocxRenderer:
    """Exports a custom template. Stateless - safe to share/reuse."""

    def render(
        self,
        template: DocumentTemplate,
        context: TemplateContext,
        output_path: str,
        screenshots: Optional[List[ScreenshotInfo]] = None,
        section_overrides: Optional[Dict[str, bool]] = None,
    ) -> str:
        screenshots = list(screenshots or [])
        doc = docx.Document()
        self._apply_base_styles(doc)

        for section in template.enabled_sections(section_overrides):
            if section.show_title and section.title.strip():
                self._add_heading(doc, context.resolve(section.title), level=1)
            for block in section.blocks:
                self._render_block(doc, block, context, screenshots)

        return self._save(doc, output_path)

    # --- Document chrome -------------------------------------------------

    def _apply_base_styles(self, doc) -> None:
        normal = doc.styles["Normal"]
        normal.font.name = FONT_NAME
        normal.font.size = Pt(BODY_SIZE_PT)

        for level in range(1, 4):
            style = doc.styles[f"Heading {level}"]
            style.font.name = FONT_NAME
            style.font.size = Pt(HEADING_SIZES_PT.get(level, 12))
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.font.bold = True

        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

    def _save(self, doc, output_path: str) -> str:
        output_path = os.path.abspath(output_path)
        parent = os.path.dirname(output_path)
        if parent:
            os.makedirs(parent, exist_ok=True)
        doc.save(output_path)
        return output_path

    # --- Block dispatch --------------------------------------------------

    def _render_block(
        self,
        doc,
        block: TemplateBlock,
        context: TemplateContext,
        screenshots: List[ScreenshotInfo],
    ) -> None:
        kind = block.kind

        if kind == BLOCK_HEADING:
            self._add_heading(doc, context.resolve(block.text), level=block.level)
        elif kind == BLOCK_SUBHEADING:
            self._add_heading(doc, context.resolve(block.text), level=max(2, block.level))
        elif kind == BLOCK_PARAGRAPH:
            self._add_paragraph(doc, context.resolve(block.text))
        elif kind == BLOCK_BULLET_LIST:
            self._add_list(doc, [context.resolve(i) for i in block.items if i.strip()], block.ordered)
        elif kind == BLOCK_TABLE:
            self._add_table(doc, block, context)
        elif kind == BLOCK_IMAGE:
            self._add_image(doc, block.image_path, context.resolve(block.caption))
        elif kind == BLOCK_SCREENSHOTS:
            self._add_screenshots(doc, block, context, screenshots)
        elif kind == BLOCK_CODE:
            self._add_code_block(doc, context.resolve(block.text))
        elif kind == BLOCK_TWO_COLUMNS:
            self._add_two_columns(doc, [context.resolve(c) for c in block.columns])
        elif kind == BLOCK_SEPARATOR:
            self._add_divider(doc)
        # Unknown kinds are ignored: a template written by a newer DAT still
        # exports everything this version understands.

    # --- Block renderers -------------------------------------------------

    def _add_heading(self, doc, text: str, level: int = 1) -> None:
        if not text:
            return
        level = max(1, min(3, int(level)))
        heading = doc.add_heading(level=level)
        run = heading.add_run(text)
        run.font.name = FONT_NAME
        run.font.size = Pt(HEADING_SIZES_PT.get(level, 12))
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_paragraph(self, doc, text: str) -> None:
        paragraph = doc.add_paragraph()
        # Author-entered newlines become real line breaks inside one paragraph.
        for index, line in enumerate((text or "").split("\n")):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE_PT)
            run.font.color.rgb = RGBColor(0, 0, 0)
        paragraph.paragraph_format.space_after = Pt(8)

    def _add_list(self, doc, items: List[str], ordered: bool = False) -> None:
        style = "List Number" if ordered else "List Bullet"
        for item in items:
            paragraph = doc.add_paragraph(style=style)
            run = paragraph.add_run(item)
            run.font.name = FONT_NAME
            run.font.size = Pt(BODY_SIZE_PT)
            run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_table(self, doc, block: TemplateBlock, context: TemplateContext) -> None:
        rows = [[context.resolve(cell) for cell in row] for row in block.table_rows]
        headers = [context.resolve(h) for h in block.table_headers] if block.include_headers else []
        col_count = len(headers) or (len(rows[0]) if rows else 0)
        if col_count == 0:
            return

        total_rows = len(rows) + (1 if headers else 0)
        if total_rows == 0:
            return

        table = doc.add_table(rows=total_rows, cols=col_count)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Honour the template's relative column widths. Word only respects
        # explicit widths with autofit off and a fixed layout, and it wants
        # the width on every cell as well as the column.
        table.autofit = False
        self._set_fixed_layout(table)
        col_widths = [
            Inches(round(CONTENT_WIDTH_IN * fraction, 3))
            for fraction in self._column_fractions(block, col_count)
        ]
        for index, column in enumerate(table.columns):
            column.width = col_widths[index]

        row_offset = 0
        if headers:
            for col in range(col_count):
                value = headers[col] if col < len(headers) else ""
                self._fill_cell(table.rows[0].cells[col], value, bold=True, width=col_widths[col])
            row_offset = 1

        for r, row in enumerate(rows):
            for col in range(col_count):
                value = row[col] if col < len(row) else ""
                self._fill_cell(table.rows[r + row_offset].cells[col], value, width=col_widths[col])

        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _column_fractions(self, block: TemplateBlock, col_count: int) -> List[float]:
        """Each column's share of the printable width, always ``col_count`` long."""
        fractions = block.col_width_fractions()
        if len(fractions) != col_count or not fractions:
            return [1.0 / col_count] * col_count
        return fractions

    def _set_fixed_layout(self, table) -> None:
        """Pin the table to a fixed layout so Word keeps our column widths
        instead of re-fitting them to the content."""
        tbl_pr = table._tbl.tblPr
        for existing in tbl_pr.findall(qn("w:tblLayout")):
            tbl_pr.remove(existing)
        tbl_pr.append(parse_xml(f'<w:tblLayout {nsdecls("w")} w:type="fixed"/>'))

    def _fill_cell(self, cell, text: str, bold: bool = False, width=None) -> None:
        if width is not None:
            cell.width = width
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.paragraph_format.left_indent = Pt(6)
        run = paragraph.add_run(str(text))
        run.font.name = FONT_NAME
        run.font.size = Pt(BODY_SIZE_PT)
        run.font.bold = bold
        run.font.color.rgb = RGBColor(0, 0, 0)

    def _add_image(self, doc, image_path: str, caption: str = "") -> None:
        if not image_path or not os.path.exists(image_path):
            return
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            paragraph.add_run().add_picture(image_path, width=Inches(self._fit_width(image_path)))
        except Exception as e:
            # A corrupt/unsupported image must not abort the whole export.
            print(f"[Warning] Skipping image {image_path}: {e}")
            return
        if caption:
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cap.add_run(caption)
            run.font.name = FONT_NAME
            run.font.size = Pt(9)
            run.font.italic = True

    def _fit_width(self, image_path: str, max_width_in: float = CONTENT_WIDTH_IN) -> float:
        """Width in inches that keeps the image inside the printable area."""
        try:
            with Image.open(image_path) as img:
                width, height = img.size
        except Exception:
            return max_width_in
        if width <= 0 or height <= 0:
            return max_width_in
        aspect = height / float(width)
        if max_width_in * aspect > MAX_IMAGE_HEIGHT_IN:
            return round(MAX_IMAGE_HEIGHT_IN / aspect, 2)
        return max_width_in

    def _add_screenshots(
        self,
        doc,
        block: TemplateBlock,
        context: TemplateContext,
        screenshots: List[ScreenshotInfo],
    ) -> None:
        if not screenshots:
            return

        heading = context.resolve(block.text) or "Screenshots"
        self._add_heading(doc, heading, level=1)

        groups = group_screenshots_by_test_case(screenshots, context.test_cases)
        for group_index, (_case_index, label, shots) in enumerate(groups):
            if label:
                paragraph = doc.add_paragraph()
                run = paragraph.add_run(label)
                run.font.name = FONT_NAME
                run.font.size = Pt(13)
                run.font.bold = False
                paragraph.paragraph_format.space_after = Pt(8)

            self._add_screenshot_grid(doc, shots)

            if group_index < len(groups) - 1 and shots:
                doc.add_page_break()

    def _add_screenshot_grid(self, doc, shots: List[ScreenshotInfo]) -> None:
        """Portrait shots pair up side-by-side; landscape shots go full width."""
        portrait_queue: List[ScreenshotInfo] = []

        def flush() -> None:
            if not portrait_queue:
                return
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for i, shot in enumerate(portrait_queue):
                paragraph = table.cell(0, i).paragraphs[0]
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    paragraph.add_run().add_picture(shot.file_path, width=Inches(MAX_PORTRAIT_WIDTH_IN))
                except Exception as e:
                    print(f"[Warning] Skipping screenshot {shot.file_path}: {e}")
            portrait_queue.clear()

        for shot in shots:
            if not os.path.exists(shot.file_path):
                continue
            try:
                with Image.open(shot.file_path) as img:
                    is_portrait = img.height > img.width
            except Exception:
                is_portrait = True

            if is_portrait:
                portrait_queue.append(shot)
                if len(portrait_queue) == 2:
                    flush()
            else:
                flush()
                self._add_image(doc, shot.file_path, shot.caption or "")

        flush()

    def _add_code_block(self, doc, code_text: str) -> None:
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        cell = table.cell(0, 0)
        self._set_cell_background(cell, "F4F5F7")

        paragraph = cell.paragraphs[0]
        for index, line in enumerate((code_text or "").split("\n")):
            if index:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.name = CODE_FONT_NAME
            run.font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _add_two_columns(self, doc, columns: List[str]) -> None:
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        half = Inches(CONTENT_WIDTH_IN / 2)
        for i in range(2):
            text = columns[i] if i < len(columns) else ""
            cell = table.cell(0, i)
            cell.width = half
            paragraph = cell.paragraphs[0]
            for index, line in enumerate(str(text).split("\n")):
                if index:
                    paragraph.add_run().add_break()
                run = paragraph.add_run(line)
                run.font.name = FONT_NAME
                run.font.size = Pt(BODY_SIZE_PT)
        doc.add_paragraph().paragraph_format.space_after = Pt(8)

    def _set_cell_background(self, cell, hex_color: str) -> None:
        tc_pr = cell._element.get_or_add_tcPr()
        tc_pr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>'))

    def _add_divider(self, doc, color_hex: str = "D0D0D0") -> None:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(6)
        paragraph.paragraph_format.space_after = Pt(12)
        border = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'<w:bottom w:val="single" w:sz="8" w:space="1" w:color="{color_hex}"/>'
            f'</w:pBdr>'
        )
        paragraph._element.get_or_add_pPr().append(border)
