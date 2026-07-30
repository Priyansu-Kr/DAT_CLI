import os
import unittest
import tempfile
import docx
from dat.renderers.docx_renderer import DocxRenderer
from dat.models.doc_request import DocRequest, ChangeSummary
from dat.models.git_info import GitInfo

class TestDocxRenderer(unittest.TestCase):
    def test_docx_renderer_generates_valid_file(self):
        renderer = DocxRenderer()
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_file = os.path.join(tmp_dir, "test_doc.docx")

            git_info = GitInfo(
                branch_name="feature/PROJ-99-nav-drawer",
                inferred_title="Navigation Drawer (PROJ-99)",
                ticket_id="PROJ-99",
                repo_name="MobileApp",
                changed_files=["app/src/Nav.kt", "app/res/layout/nav.xml"],
                raw_diff="+ val drawer = DrawerLayout()\n- val drawer = null"
            )

            summary = ChangeSummary(
                overview="Implemented navigation drawer layout.",
                key_points=["Added Nav.kt drawer handler.", "Created xml layout."],
                impact_areas=["UI Navigation"],
                test_recommendations=["Open drawer on swipe gesture."]
            )

            doc_req = DocRequest(
                title="Navigation Drawer Implementation",
                subtitle="Feature Specification",
                author="Test Developer",
                ticket_id="PROJ-99",
                git_info=git_info,
                summary=summary,
                output_path=output_file
            )

            result_path = renderer.render(doc_req)
            self.assertTrue(os.path.exists(result_path))

            document = docx.Document(result_path)
            text_content = "\n".join([p.text for p in document.paragraphs])
            self.assertIn("Navigation Drawer Implementation", text_content)
            # AI Summary section was removed entirely - it must never render,
            # even when a summary with overview text is present.
            self.assertNotIn("AI Summary", text_content)
            self.assertNotIn("Implemented navigation drawer layout.", text_content)

    def test_docx_renderer_respects_section_toggles(self):
        renderer = DocxRenderer()
        with tempfile.TemporaryDirectory() as tmp_dir:
            output_file = os.path.join(tmp_dir, "toggled_doc.docx")

            summary = ChangeSummary(
                overview="Implemented navigation drawer layout.",
                key_points=["Added Nav.kt drawer handler."],
                impact_areas=["UI Navigation"],
                test_cases=["Open drawer via swipe gesture"],
            )

            doc_req = DocRequest(
                title="Navigation Drawer Implementation",
                author="Test Developer",
                ticket_id="PROJ-99",
                summary=summary,
                output_path=output_file,
                sections={
                    "header": False,
                    "metadata_table": False,
                    "changes_done": True,
                    "test_cases_table": False,
                    "screenshots": False,
                },
            )

            result_path = renderer.render(doc_req)
            document = docx.Document(result_path)
            text_content = "\n".join([p.text for p in document.paragraphs])

            self.assertNotIn("Navigation Drawer Implementation", text_content)
            self.assertNotIn("Task Detail", text_content)
            self.assertIn("Changes Done", text_content)
            self.assertIn("Added Nav.kt drawer handler.", text_content)
            self.assertEqual(len(document.tables), 0)

if __name__ == "__main__":
    unittest.main()
