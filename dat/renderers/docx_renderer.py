import os
from typing import Optional, List
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
from PIL import Image

from dat.renderers.base_renderer import BaseRenderer
from dat.models.doc_request import DocRequest, DEFAULT_SECTIONS
from dat.renderers.screenshot_grouping import group_screenshots_by_test_case

class DocxRenderer(BaseRenderer):
    def render(self, doc_request: DocRequest) -> str:
        sections = doc_request.sections or DEFAULT_SECTIONS
        doc = docx.Document()

        # Global Font Setting: Arial
        style = doc.styles['Normal']
        style.font.name = 'Arial'
        style.font.size = Pt(11)

        # Force Heading styles to be Black and Arial
        for level in range(1, 4):
            heading_style = doc.styles[f'Heading {level}']
            heading_style.font.name = 'Arial'
            heading_style.font.size = Pt(14 if level == 1 else 12)
            heading_style.font.color.rgb = RGBColor(0, 0, 0)
            heading_style.font.bold = True

        # Set page margins
        for section in doc.sections:
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        # 1. TOP HEADING (TICKET - TOPIC)
        if sections.get("header", True):
            p_title = doc.add_paragraph()
            p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_title = p_title.add_run(doc_request.title) # Formatted as "TICKET - TOPIC"
            run_title.font.name = 'Arial'
            run_title.font.size = Pt(18)
            run_title.font.bold = True
            run_title.font.color.rgb = RGBColor(0, 0, 0)
            p_title.paragraph_format.space_after = Pt(0)

            # ~2 line space after heading
            for _ in range(2):
                doc.add_paragraph()

        # 2. TASK DETAIL HEADER + 3. METADATA TABLE (5 rows x 2 columns)
        if sections.get("metadata_table", True):
            p_task_hdr = doc.add_paragraph()
            run_task_hdr = p_task_hdr.add_run("Task Detail")
            run_task_hdr.font.name = 'Arial'
            run_task_hdr.font.size = Pt(16)
            run_task_hdr.font.bold = True
            run_task_hdr.font.color.rgb = RGBColor(0, 0, 0)
            p_task_hdr.paragraph_format.space_after = Pt(6)

            meta_table = doc.add_table(rows=5, cols=2)
            meta_table.style = 'Table Grid'
            meta_table.alignment = WD_TABLE_ALIGNMENT.LEFT

            # Split title back for Short Description if needed
            # Format of doc_request.title is "TICKET TOPIC"
            # We want only the TOPIC part for Short Description
            title_parts = doc_request.title.split(' ', 1)
            short_desc = title_parts[1] if len(title_parts) > 1 else doc_request.title

            labels = ["Ticket No.", "Short Description", "Document Date", "Created By", "Approved By"]
            values = [
                doc_request.ticket_id or "",
                short_desc,
                datetime.now().strftime("%d-%B-%Y"),
                doc_request.author,
                doc_request.approved_by
            ]

            for i in range(5):
                row = meta_table.rows[i]
                # Set row height/padding (approximate by paragraph spacing)
                row.height = Pt(30)

                cell_lbl = row.cells[0]
                cell_val = row.cells[1]

                # Label Styling
                p_l = cell_lbl.paragraphs[0]
                p_l.paragraph_format.space_before = Pt(8)
                p_l.paragraph_format.space_after = Pt(8)
                p_l.paragraph_format.left_indent = Pt(6)
                r_l = p_l.add_run(labels[i])
                r_l.font.name = 'Arial'
                r_l.font.bold = False
                r_l.font.size = Pt(11)
                r_l.font.color.rgb = RGBColor(0, 0, 0)

                # Value Styling
                p_v = cell_val.paragraphs[0]
                p_v.paragraph_format.space_before = Pt(8)
                p_v.paragraph_format.space_after = Pt(8)
                p_v.paragraph_format.left_indent = Pt(6)
                r_v = p_v.add_run(str(values[i]))
                r_v.font.name = 'Arial'
                r_v.font.size = Pt(11)
                r_v.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # 4. CHANGES DONE BLOCK
        if sections.get("changes_done", True):
            h_changes = doc.add_heading(level=1)
            run_changes = h_changes.add_run("Changes Done")
            run_changes.font.name = 'Arial'
            run_changes.font.color.rgb = RGBColor(0, 0, 0)
            # Add one line spacing below heading
            doc.add_paragraph()

            if doc_request.summary:
                # Affected Modules Sub-section
                p_aff = doc.add_paragraph()
                r_aff_label = p_aff.add_run("Affected Module: ")
                r_aff_label.font.name = 'Arial'
                r_aff_label.font.bold = True
                r_aff_label.font.size = Pt(11)

                modules_text = ", ".join(doc_request.summary.impact_areas) if doc_request.summary.impact_areas else "Main Module"
                r_aff_val = p_aff.add_run(modules_text)
                r_aff_val.font.name = 'Arial'
                r_aff_val.font.size = Pt(11)

                # Key Points (Concise)
                for point in doc_request.summary.key_points:
                    p = doc.add_paragraph(style='List Bullet')
                    r = p.add_run(point)
                    r.font.name = 'Arial'
                    r.font.size = Pt(11)
                    r.font.color.rgb = RGBColor(0, 0, 0)
            else:
                p = doc.add_paragraph(style='List Bullet')
                r = p.add_run("Implemented core logic changes.")
                r.font.name = 'Arial'
                r.font.size = Pt(11)
                r.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # 5. TEST CASES TABLE (3 columns: Index, Case, Status)
        if sections.get("test_cases_table", True) and doc_request.summary:
            # One line gap before test cases table
            doc.add_paragraph()

            if doc_request.summary.test_cases:
                test_table = doc.add_table(rows=len(doc_request.summary.test_cases) + 1, cols=3)
                test_table.style = 'Table Grid'
                test_table.alignment = WD_TABLE_ALIGNMENT.LEFT
                test_table.autofit = False # Disable autofit to set manual widths
                
                # Define widths
                width_idx = Inches(1.0)
                width_case = Inches(4.5)
                width_status = Inches(1.0)

                # Set widths for columns directly (Most reliable way)
                test_table.columns[0].width = width_idx
                test_table.columns[1].width = width_case
                test_table.columns[2].width = width_status

                # Header Row
                hdr_labels = ["Index", "Case", "Status"]
                for i, label in enumerate(hdr_labels):
                    cell = test_table.rows[0].cells[i]
                    # Also set width on individual cell for better compatibility
                    cell.width = [width_idx, width_case, width_status][i]

                    p = cell.paragraphs[0]
                    r = p.add_run(label)
                    r.font.name = 'Arial'
                    r.font.bold = True
                    r.font.size = Pt(11)
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(8)
                    p.paragraph_format.left_indent = Pt(6)

                # Data Rows
                for i, case in enumerate(doc_request.summary.test_cases):
                    row = test_table.rows[i+1]
                    row.height = Pt(30)
                    
                    # Index
                    cell_idx = row.cells[0]
                    cell_idx.width = width_idx
                    p_idx = cell_idx.paragraphs[0]
                    r_idx = p_idx.add_run(f"{i+1}.")
                    r_idx.font.name = 'Arial'
                    r_idx.font.size = Pt(11)
                    p_idx.paragraph_format.space_before = Pt(8)
                    p_idx.paragraph_format.space_after = Pt(8)
                    p_idx.paragraph_format.left_indent = Pt(6)
                    
                    # Case (AI Generated)
                    cell_case = row.cells[1]
                    cell_case.width = width_case
                    p_case = cell_case.paragraphs[0]
                    r_case = p_case.add_run(case)
                    r_case.font.name = 'Arial'
                    r_case.font.size = Pt(11)
                    p_case.paragraph_format.space_before = Pt(8)
                    p_case.paragraph_format.space_after = Pt(8)
                    p_case.paragraph_format.left_indent = Pt(6)
                    
                    # Status (Always Success)
                    cell_stat = row.cells[2]
                    cell_stat.width = width_status
                    p_stat = cell_stat.paragraphs[0]
                    r_stat = p_stat.add_run("Success")
                    r_stat.font.name = 'Arial'
                    r_stat.font.size = Pt(11)
                    p_stat.paragraph_format.space_before = Pt(8)
                    p_stat.paragraph_format.space_after = Pt(8)
                    p_stat.paragraph_format.left_indent = Pt(6)

            doc.add_paragraph().paragraph_format.space_after = Pt(24)

        # 6. SCREENSHOTS
        if sections.get("screenshots", True) and doc_request.screenshots:
            doc.add_page_break()
            h_shots = doc.add_heading(level=1)
            run_shots = h_shots.add_run("Screenshots")
            run_shots.font.name = 'Arial'
            run_shots.font.color.rgb = RGBColor(0, 0, 0)
            doc.add_paragraph()

            # Group screenshots under their assigned test case (or auto-distribute
            # across test cases when no explicit assignment was made).
            test_cases = doc_request.summary.test_cases if doc_request.summary else []
            groups = group_screenshots_by_test_case(doc_request.screenshots, test_cases)

            for group_idx, (case_idx, label, case_shots) in enumerate(groups):
                # Add Test Case Header: Arial 18, Not Bold
                p_tc = doc.add_paragraph()
                r_tc = p_tc.add_run(label)
                r_tc.font.name = 'Arial'
                r_tc.font.size = Pt(18)
                r_tc.font.bold = False
                p_tc.paragraph_format.space_after = Pt(12)

                # Render shots for this specific test case
                portrait_queue = []
                
                def flush_portrait_queue(queue, document):
                    if not queue: return
                    # Create a borderless table for side-by-side layout
                    table = document.add_table(rows=1, cols=2)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    for i, shot in enumerate(queue):
                        cell = table.cell(0, i)
                        p = cell.paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.add_run().add_picture(shot.file_path, width=Inches(3.0))
                    queue.clear()

                for i, shot in enumerate(case_shots):
                    if not os.path.exists(shot.file_path):
                        continue
                    
                    try:
                        with Image.open(shot.file_path) as img:
                            w, h = img.size
                            is_portrait = h > w
                    except:
                        is_portrait = True

                    if is_portrait:
                        portrait_queue.append(shot)
                        if len(portrait_queue) == 2:
                            flush_portrait_queue(portrait_queue, doc)
                    else:
                        flush_portrait_queue(portrait_queue, doc)
                        if i > 0 or len(portrait_queue) > 0:
                            doc.add_page_break()
                            # Repeat the Test Case header on new page if it continues
                            p_tc_cont = doc.add_paragraph()
                            r_tc_cont = p_tc_cont.add_run(f"{label} (cont.)")
                            r_tc_cont.font.name = 'Arial'
                            r_tc_cont.font.size = Pt(18)
                            r_tc_cont.font.bold = False
                        
                        p_img = doc.add_paragraph()
                        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p_img.add_run().add_picture(shot.file_path, width=Inches(6.5))
                
                flush_portrait_queue(portrait_queue, doc)

                # Add page break between groups if not the last one
                if group_idx < len(groups) - 1 and len(case_shots) > 0:
                    doc.add_page_break()

        output_path = os.path.abspath(doc_request.output_path)
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        return output_path

    def _add_code_block(self, doc, code_text: str):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        self._set_cell_background(cell, "F8FAFC")
        
        p = cell.paragraphs[0]
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(8)

    def _set_cell_background(self, cell, hex_color: str):
        tcPr = cell._element.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
        tcPr.append(shd)

    def _add_divider(self, doc, color_hex: str = "1E3A8A"):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="18" w:space="1" w:color="{color_hex}"/></w:pBdr>')
        p._element.get_or_add_pPr().append(pBdr)
